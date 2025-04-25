from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from app.utils.file_utils import delete_file

class DocxService:
    def __init__(
        self, 
        template_path: str, 
    ):
        print("Initializing...")
        self.doc = Document(template_path)

    # Replace a marker with an image in the document.
    def replace_marker_with_image(self, marker: str, image_path: str):
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if marker in cell.text:
                        cell.text = ""
                        cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = cell.paragraphs[0].add_run()
                        run.add_picture(image_path, width=Inches(1.5))

    def replace_text(self, replacements: dict):
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        if key in cell.text:
                            print("Replacing:", key, "with", value)
                            print("Old text:", cell.text)
                            cell.text = cell.text.replace(key, value)
                            print("New text:", cell.text)
                            
    # This method replaces text in the document while preserving the style of the text.
    def replace_text_preserve_style(self, replacements: dict):
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for key, value in replacements.items():
                                print("Checking:", key, "in", run.text)
                                if key in run.text:
                                    print("Replacing:", key, "with", value)
                                    run.text = run.text.replace(key, value)
                                    
    # This method replaces text in the header of the document.
    def replace_header_text(self, replacements: dict):
        for section in self.doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    for key, value in replacements.items():
                        if key in run.text:
                            run.text = run.text.replace(key, value)

    # This method saves the document to a specified output path.
    def save(self, output_path: str):
        self.doc.save(output_path)
        
    # This method show the formatting options of the document.
    def formating_options(self):
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    print(cell.text)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            print(run.text)
                            
